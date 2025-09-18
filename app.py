import locale
from flask import Flask, render_template, abort, url_for, request, redirect, flash, jsonify, send_file
import datetime
import os
from werkzeug.utils import secure_filename
from ocr_processor import processar_nf
from pypdf import PdfWriter
from openpyxl import Workbook
from openpyxl.styles import Font
from io import BytesIO
from docx import Document
import calendar
import boto3
from botocore.exceptions import NoCredentialsError
from database import get_db_connection, criar_tabelas
from babel.numbers import format_decimal
import psycopg2.extras

# --- CONFIGURAÇÃO DE CAMINHO DINÂMICO ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# --- CONFIGURAÇÃO DE LOCAL E FLASK ---
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    print("AVISO: Locale 'pt_BR.UTF-8' não encontrado.")

try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    print("AVISO: Locale 'pt_BR.UTF-8' não encontrado.")

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SECRET_KEY'] = 'uma-chave-secreta-muito-dificil'
UPLOAD_FOLDER = '/tmp/uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Cria as tabelas da base de dados no arranque, se necessário
try:
    criar_tabelas()
except Exception as e:
    print(f"Ocorreu um erro ao inicializar a base de dados: {e}")

# --- CONFIGURAÇÃO DO CLOUDFLARE R2 ---
CLOUDFLARE_ACCOUNT_ID = os.environ.get('CLOUDFLARE_ACCOUNT_ID')
AWS_ACCESS_KEY_ID = os.environ.get('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY')
BUCKET_NAME = os.environ.get('BUCKET_NAME')

s3_client = None
if CLOUDFLARE_ACCOUNT_ID:
    R2_ENDPOINT_URL = f'https://{CLOUDFLARE_ACCOUNT_ID}.r2.cloudflarestorage.com'
    s3_client = boto3.client(
        's3',
        endpoint_url=R2_ENDPOINT_URL,
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        region_name='auto'
    )

@app.template_filter('formatar_valor')
def formatar_valor(value):
    if value is None: return "0,00"
    try:
        return format_decimal(value, locale='pt_BR')
    except (ValueError, TypeError):
        return value

def get_db():
    conn = get_db_connection()
    # Usar um cursor que retorna dicionários
    return conn, conn.cursor(cursor_factory=psycopg2.extras.DictCursor)

# --- ROTA PRINCIPAL ---
@app.route('/')
def index():
    conn, db = get_db()
    db.execute('SELECT * FROM professores ORDER BY nome') # 1. Executar a consulta
    professores_db = db.fetchall()                      # 2. Obter os resultados
    db.close()
    conn.close()
    professores_cultura = [p for p in professores_db if p['categoria'] == 'Cultura']
    professores_esporte = [p for p in professores_db if p['categoria'] == 'Esporte']
    return render_template('index.html', cultura=professores_cultura, esporte=professores_esporte)

# --- ROTAS DE GESTÃO DE PROFESSORES ---
@app.route('/professor/adicionar', methods=['GET', 'POST'])
def adicionar_professor():
    if request.method == 'POST':
        nome = request.form['nome']
        categoria = request.form['categoria']
        cpf = request.form['cpf'] or None
        cnpj = request.form['cnpj'] or None
        dados_bancarios = request.form['dados_bancarios'] or None
        conn = get_db_connection()  
        db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        db.execute(
        'INSERT INTO professores (nome, categoria, cpf, cnpj, dados_bancarios) VALUES (%s, %s, %s, %s, %s)',
        (nome, categoria, cpf, cnpj, dados_bancarios)
        )
        conn.commit()
        db.close()
        conn.close()
        flash('Professor adicionado com sucesso!', 'success')
        return redirect(url_for('index'))
    return render_template('adicionar_professor.html')

@app.route('/professor/editar/<int:professor_id>', methods=['GET', 'POST'])
def editar_professor(professor_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('SELECT * FROM professores WHERE id = %s', (professor_id,))
    professor = db.fetchone()
    if request.method == 'POST':
        nome = request.form['nome']
        categoria = request.form['categoria']
        cpf = request.form['cpf'] or None
        cnpj = request.form['cnpj'] or None
        dados_bancarios = request.form['dados_bancarios'] or None
        db.execute(
        'UPDATE professores SET nome = %s, categoria = %s, cpf = %s, cnpj = %s, dados_bancarios = %s WHERE id = %s',
        (nome, categoria, cpf, cnpj, dados_bancarios, professor_id)
        )
        conn.commit()
        db.close()
        conn.close()
        flash('Dados do professor atualizados com sucesso!', 'success')
        return redirect(url_for('detalhes_professor', professor_id=professor_id))
    return render_template('editar_professor.html', professor=professor)

@app.route('/professor/deletar/<int:professor_id>', methods=['POST'])
def deletar_professor(professor_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('DELETE FROM professores WHERE id = %s', (professor_id,))
    conn.commit()
    db.close()
    conn.close()
    flash('Professor apagado com sucesso.', 'success')
    return redirect(url_for('index'))

# --- ROTAS DE DETALHES, UPLOAD E VISUALIZAÇÃO ---
@app.route('/professor/<int:professor_id>')
def detalhes_professor(professor_id):
    conn, db = get_db() 
    db.execute('SELECT * FROM professores WHERE id = %s', (professor_id,))
    professor = db.fetchone()
    db.close()
    conn.close()
    if professor is None: abort(404)
    anos = [datetime.datetime.now().year + i for i in range(3)]
    meses = [(m, datetime.date(2000, m, 1).strftime('%B').capitalize()) for m in range(1, 13)]
    return render_template('professor_detalhes.html', professor=professor, meses=meses, anos=anos)

@app.route('/professor/<int:professor_id>/<int:ano>/<int:mes>', methods=['GET', 'POST'])
def mes_detalhes(professor_id, ano, mes):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('SELECT * FROM professores WHERE id = %s', (professor_id,))
    professor = db.fetchone()
    if professor is None: abort(404)
    nome_mes = datetime.date(ano, mes, 1).strftime('%B').capitalize()

    if request.method == 'POST':
        if not s3_client:
            flash("Configuração de armazenamento na nuvem não encontrada.", "error")
            return redirect(request.url)
            
        arquivos = {'NF': request.files.get('nota_fiscal'), 'Relatorio': request.files.get('relatorio'), 'Chamada': request.files.get('chamada')}
        for tipo, arquivo in arquivos.items():
            if arquivo and arquivo.filename != '':
                filename_seguro = secure_filename(arquivo.filename)
                
                # CÓDIGO CORRIGIDO
                caminho_temporario = os.path.join(app.config['UPLOAD_FOLDER'], filename_seguro)
                arquivo.save(caminho_temporario)

                # Processa o arquivo do caminho salvo
                if tipo == 'NF':
                    dados_ocr = processar_nf(caminho_temporario)
                    print(f"DEBUG: Dados extraídos do OCR para a NF: {dados_ocr}")
                    nf_numero = dados_ocr.get('numero')
                    nf_data = dados_ocr.get('data')
                    nf_valor = dados_ocr.get('valor')
                else:
                    dados_ocr = {}
                    nf_numero = None
                    nf_data = None
                    nf_valor = None
                
                # Faz o upload para o R2 e remove o arquivo temporário
                nome_final_r2 = f"{professor_id}/{ano}/{mes}/{tipo}_{filename_seguro}"
                try:
                    s3_client.upload_file(caminho_temporario, BUCKET_NAME, nome_final_r2)
                    os.remove(caminho_temporario)
                    db.execute(
                    'INSERT INTO documentos (professor_id, mes, ano, tipo_documento, caminho_arquivo, nf_numero, nf_data, nf_valor) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)',
                    (professor_id, mes, ano, tipo, nome_final_r2, nf_numero, nf_data, nf_valor)
                    )
                except Exception as e:
                    flash(f"Erro ao fazer upload: {e}", "error")
        
        conn.commit()
        db.close()
        conn.close()
        flash('Documentos enviados com sucesso!', 'success')
        return redirect(url_for('mes_detalhes', professor_id=professor_id, ano=ano, mes=mes))

    db.execute('SELECT * FROM documentos WHERE professor_id = %s AND mes = %s AND ano = %s', (professor_id, mes, ano))
    docs_db = db.fetchall()
    documentos = {doc['tipo_documento']: doc for doc in docs_db}
    db.close()
    return render_template('mes_detalhes.html', professor=professor, mes_numero=mes, nome_mes=nome_mes, ano=ano, documentos=documentos)

@app.route('/view/<path:filename>')
def view_file(filename):
    try:
        url = s3_client.generate_presigned_url('get_object', Params={'Bucket': BUCKET_NAME, 'Key': filename}, ExpiresIn=300)
        return redirect(url)
    except Exception as e:
        abort(404, description=f"Erro ao aceder ao ficheiro: {e}")

@app.route('/documento/deletar/<int:doc_id>', methods=['POST'])
def deletar_documento(doc_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('SELECT * FROM documentos WHERE id = %s', (doc_id,))
    doc = db.fetchone()
    if doc:
        try:
            s3_client.delete_object(Bucket=BUCKET_NAME, Key=doc['caminho_arquivo'])
        except Exception as e:
            flash(f'Erro ao apagar o arquivo da nuvem: {e}', 'error')
        db.execute('DELETE FROM documentos WHERE id = %s', (doc_id,))
        conn.commit()
        flash('Documento apagado com sucesso.', 'success')
        professor_id = doc['professor_id']
        ano = doc['ano']
        mes = doc['mes']
        db.close()
        conn.close()
        return redirect(url_for('mes_detalhes', professor_id=professor_id, ano=ano, mes=mes))
    db.close()
    conn.close()
    return redirect(url_for('index'))

@app.route('/ferramentas')
def ferramentas_pdf():
    return render_template('ferramentas.html')

@app.route('/ferramentas/juntar-pdf', methods=['GET', 'POST'])
def juntar_pdf():
    if request.method == 'POST':
        files = request.files.getlist('arquivos_pdf')
        if len(files) < 2:
            flash('Selecione pelo menos dois arquivos.', 'error')
            return redirect(request.url)
        merger = PdfWriter()
        for pdf in files:
            merger.append(pdf)
        output_io = BytesIO()
        merger.write(output_io)
        merger.close()
        output_io.seek(0)
        return send_file(output_io, as_attachment=True, download_name='documento_juntado.pdf', mimetype='application/pdf')
    return render_template('juntar_pdf.html')

@app.route('/ferramentas/converter-word', methods=['GET', 'POST'])
def converter_word():
    if request.method == 'POST':
        word_file = request.files.get('arquivo_word')

        if not word_file or word_file.filename == '':
            flash('Por favor, selecione um arquivo Word.', 'error')
            return redirect(request.url)

        if word_file and (word_file.filename.lower().endswith('.docx') or word_file.filename.lower().endswith('.doc')):
            filename_seguro = secure_filename(word_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_seguro)
            word_file.save(input_path)

            output_filename = os.path.splitext(filename_seguro)[0] + '.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            try:
                convert(input_path, output_path)
                os.remove(input_path)
                return send_from_directory(app.config['UPLOAD_FOLDER'], output_filename, as_attachment=True)
            except Exception as e:
                flash(f'Ocorreu um erro durante a conversão: {e}', 'error')
                if os.path.exists(input_path):
                    os.remove(input_path)
                return redirect(request.url)
        else:
            flash('Formato de ficheiro inválido. Por favor, envie um ficheiro .doc ou .docx.', 'error')
            return redirect(request.url)

    return render_template('converter_word.html')

@app.route('/controle-gastos')
def controle_gastos_index():
    # Esta página agora só mostra a escolha entre Cultura e Esporte
    return render_template('controle_gastos_index.html')

@app.route('/controle-gastos/<categoria>')
def gastos_por_categoria(categoria):
    if categoria not in ['Cultura', 'Esporte']:
        abort(404)

    ano_atual = datetime.datetime.now().year
    anos = [ano_atual + i for i in range(3)]
    parcelas = {
        1: "Fevereiro, Março e Abril",
        2: "Maio e Junho",
        3: "Julho e Agosto",
        4: "Setembro e Outubro",
        5: "Novembro e Dezembro"
    }
    return render_template('gastos_por_categoria.html', categoria=categoria, anos=anos, parcelas=parcelas)


@app.route('/controle-gastos/<categoria>/<int:ano>/<int:parcela>', methods=['GET', 'POST'])
def parcela_gastos(categoria, ano, parcela):
    conn, db = get_db() # Desempacotar aqui

    if request.method == 'POST':
        # --- LÓGICA PARA SALVAR O VALOR INICIAL ---
        valor_inicial_str = request.form.get('valor_inicial').replace('R$', '').strip().replace('.', '').replace(',', '.')
        valor_inicial_float = float(valor_inicial_str)

        # Verifica se já existe um registo para esta parcela
        conn = get_db_connection()
        db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        db.execute("""
        INSERT INTO parcelas (categoria, ano, parcela,  valor_inicial) 
        VALUES (%s, %s, %s, %s)
        ON CONFLICT (categoria, ano, parcela) 
        DO UPDATE SET valor_inicial = EXCLUDED.valor_inicial
        """, (categoria, ano, parcela, valor_inicial_float)
        )


        # --- Lógica para salvar os gastos
        db.execute('DELETE FROM gastos WHERE categoria = %s AND ano = %s AND parcela = %s', (categoria, ano, parcela))
        descricoes = request.form.getlist('descricao[]')
        valores = request.form.getlist('valor[]')

        for i in range(len(descricoes)):
            if descricoes[i] and valores[i]:
                try:
                    valor_float = float(valores[i].replace('.', '').replace(',', '.'))
                    db.execute(
                    'INSERT INTO gastos (categoria, ano, parcela, descricao, valor) VALUES (%s, %s, %s, %s, %s)',
                    (categoria, ano, parcela, descricoes[i], valor_float)
                    )
                except ValueError:
                    flash(f'Valor inválido "{valores[i]}" ignorado.', 'error')
        
        conn.commit()
        db.close()
        conn.close()
        flash('Gastos salvos com sucesso!', 'success')
        return redirect(url_for('parcela_gastos', categoria=categoria, ano=ano, parcela=parcela))

    # --- LÓGICA GET ATUALIZADA PARA LER O VALOR INICIAL DO BANCO DE DADOS ---
    db.execute(
    'SELECT valor_inicial FROM parcelas WHERE categoria = %s AND ano = %s AND parcela = %s',
    (categoria, ano, parcela)
    )
    parcela_info = db.fetchone()

    if parcela_info:
        valor_inicial = parcela_info['valor_inicial']
    else:
        valores_padrao = {"Cultura": 72524.62, "Esporte": 31389.00}
        valor_inicial = valores_padrao.get(categoria, 0)

    db.execute(
    'SELECT * FROM gastos WHERE categoria = %s AND ano = %s AND parcela = %s ORDER BY id',
    (categoria, ano, parcela)
    )
    gastos = db.fetchall()

    total_gasto = sum(g['valor'] for g in gastos)
    saldo = valor_inicial - total_gasto

    return render_template('parcela_gastos.html', categoria=categoria, ano=ano, parcela=parcela,
                           valor_inicial=valor_inicial, gastos=gastos, total_gasto=total_gasto, saldo=saldo)


@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)



@app.route('/gastos/deletar/<int:gasto_id>', methods=['POST'])
def deletar_gasto(gasto_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('SELECT * FROM gastos WHERE id = %s', (gasto_id,))
    gasto = db.fetchone()
    
    if gasto:
        db.execute('DELETE FROM gastos WHERE id = %s', (gasto_id,))
        conn.commit()
        flash('Gasto apagado com sucesso.', 'success')
        categoria = gasto['categoria']
        ano = gasto['ano']
        parcela = gasto['parcela']
        db.close()
        conn.close()
        return redirect(url_for('parcela_gastos', categoria=categoria, ano=ano, parcela=parcela))
    
    db.close()
    conn.close()
    flash('Gasto não encontrado.', 'error')
    return redirect(url_for('controle_gastos_index'))

@app.route('/emprestimos')
def emprestimos():
    conn, db = get_db()
    # Pega todos os empréstimos, os mais recentes primeiro
    db.execute('SELECT * FROM emprestimos ORDER BY data_retirada DESC')
    lista_emprestimos = db.fetchall()
    db.close()
    conn.close()
    return render_template('emprestimos.html', emprestimos=lista_emprestimos)

@app.route('/emprestimos/adicionar', methods=['GET', 'POST'])
def adicionar_emprestimo():
    if request.method == 'POST':
        data_retirada = request.form['data_retirada']
        item = request.form['item']
        responsavel = request.form['responsavel']
        observacoes = request.form['observacoes']

        if not data_retirada or not item or not responsavel:
            flash('Data de Retirada, Item e Responsável são campos obrigatórios.', 'error')
        else:
            conn = get_db_connection()
            db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            db.execute(
                'INSERT INTO emprestimos (data_retirada, item, responsavel, observacoes, data_devolucao) VALUES (%s, %s, %s, %s, NULL)',
                (data_retirada, item, responsavel, observacoes)
            )
            conn.commit()
            db.close()
            conn.close()
            flash('Empréstimo registado com sucesso!', 'success')
            return redirect(url_for('emprestimos'))
    
    data_hoje = datetime.datetime.now().strftime('%Y-%m-%d')
    return render_template('formulario_emprestimo.html', acao="Adicionar", data_hoje=data_hoje)

@app.route('/emprestimos/editar/<int:emprestimo_id>', methods=['GET', 'POST'])
def editar_emprestimo(emprestimo_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('SELECT * FROM emprestimos WHERE id = %s', (emprestimo_id,))
    emprestimo = db.fetchone()
    db.close()

    if emprestimo is None:
        abort(404)

    if request.method == 'POST':
        data_retirada = request.form['data_retirada']
        item = request.form['item']
        responsavel = request.form['responsavel']
        data_devolucao = request.form.get('data_devolucao') # .get para ser seguro caso não exista
        observacoes = request.form['observacoes']

        if not data_devolucao: # Se o campo for deixado vazio, guarda como NULL
            data_devolucao = None

        if not data_retirada or not item or not responsavel:
            flash('Data de Retirada, Item e Responsável são campos obrigatórios.', 'error')
        else:
            db = get_db()
            db.execute(
                'UPDATE emprestimos SET data_retirada = %s, item = %s, responsavel = %s, data_devolucao = %s, observacoes = %s WHERE id = %s',
                (data_retirada, item, responsavel, data_devolucao, observacoes, emprestimo_id)
            )
            conn.commit() 
            flash('Empréstimo atualizado com sucesso!', 'success')
            
            db.close()
            conn.close()
            return redirect(url_for('emprestimos'))

    return render_template('formulario_emprestimo.html', acao="Editar", emprestimo=emprestimo)

@app.route('/emprestimos/deletar/<int:emprestimo_id>', methods=['POST'])
def deletar_emprestimo(emprestimo_id):
    conn = get_db_connection()
    db = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    db.execute('DELETE FROM emprestimos WHERE id = %s', (emprestimo_id,))
    conn.commit()
    flash('Registo de empréstimo apagado com sucesso.', 'success')
    db.close()
    conn.close()
    return redirect(url_for('emprestimos'))

# --- ROTA PARA A PÁGINA DE RELATÓRIOS
@app.route('/relatorio', methods=['GET', 'POST'])
def relatorio():
    if request.method == 'POST':
        tipo_periodo = request.form.get('tipo_periodo')
        professor_id = request.form.get('professor_id')
        ano = int(request.form.get('ano'))
        formato = request.form.get('formato')
        mes = request.form.get('mes')

        conn, db = get_db()
        query = """
            SELECT p.nome, d.mes, d.nf_numero, d.nf_data, d.nf_valor
            FROM documentos d
            JOIN professores p ON d.professor_id = p.id
            WHERE d.tipo_documento = 'NF' AND d.ano = %s
        """
        params = [ano]
        if tipo_periodo == 'mensal' and mes:
            query += " AND d.mes = %s"
            params.append(int(mes))
        if professor_id != 'todos':
            query += " AND p.id = %s"
            params.append(int(professor_id))
        query += " ORDER BY p.nome, d.mes"
        db.execute(query, tuple(params))
        resultados = db.fetchall()
        db.close()
        conn.close()

        if not resultados:
            flash('Nenhum dado encontrado para os filtros selecionados.', 'error')
            return redirect(url_for('relatorio'))

        # --- Geração de Ficheiros ---
        if formato == 'excel':
            wb = Workbook()
            ws = wb.active
            ws.title = "Relatório de Pagamentos"
            headers = ["Professor", "Mês", "Nº da Nota", "Data da Nota", "Valor (R$)"]
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True)
            total_geral = 0
            for linha in resultados:
                valor = linha['nf_valor'] or 0
                ws.append([linha['nome'], str(linha['mes']), linha['nf_numero'], linha['nf_data'], valor])
                total_geral += valor
            ws.append([])
            ws.append(["", "", "", "TOTAL GERAL:", total_geral])
            total_cell = ws.cell(row=ws.max_row, column=4)
            total_cell.font = Font(bold=True)
            valor_total_cell = ws.cell(row=ws.max_row, column=5)
            valor_total_cell.font = Font(bold=True)
            valor_total_cell.number_format = '"R$" #,##0.00'
            virtual_workbook = BytesIO()
            wb.save(virtual_workbook)
            virtual_workbook.seek(0)
            nome_ficheiro = f"relatorio_{ano}_{mes if tipo_periodo == 'mensal' else 'anual'}.xlsx"
            return send_file(
                virtual_workbook,
                as_attachment=True,
                download_name=nome_ficheiro,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        
        # --- NOVO BLOCO PARA GERAR O FICHEIRO WORD ---
        elif formato == 'word':
            document = Document()
            document.add_heading('Relatório de Pagamentos', level=1)

            nome_professor = "Todos os Professores"
            if professor_id != 'todos':
                db = get_db()
                db.execute('SELECT nome FROM professores WHERE id = %s', (professor_id,))
                prof = db.fetchone()
                db.close()
                if prof:
                    nome_professor = prof['nome']
            
            periodo_str = f"Ano: {ano}"
            if tipo_periodo == 'mensal':
                nome_mes = datetime.date(2000, int(mes), 1).strftime('%B').capitalize()
                periodo_str = f"Período: {nome_mes} de {ano}"
            
            document.add_paragraph(f"Professor: {nome_professor}")
            document.add_paragraph(periodo_str)
            document.add_paragraph()

            # Criar a tabela
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Professor", "Mês", "Nº da Nota", "Data da Nota", "Valor (R$)"]
            for i, header_text in enumerate(headers):
                hdr_cells[i].text = header_text
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Preencher a tabela com os dados
            total_geral = 0
            for linha in resultados:
                row_cells = table.add_row().cells
                valor = linha['nf_valor'] or 0
                row_cells[0].text = linha['nome']
                row_cells[1].text = str(linha['mes'])
                row_cells[2].text = str(linha['nf_numero'] or 'N/A')
                row_cells[3].text = str(linha['nf_data'] or 'N/A')
                row_cells[4].text = f"R$ {formatar_valor(valor)}"
                total_geral += valor
            
            # Adicionar linha de total
            document.add_paragraph() # Espaçamento
            p = document.add_paragraph()
            p.add_run('TOTAL GERAL: ').bold = True
            p.add_run(f"R$ {formatar_valor(total_geral)}").bold = True

            # Salvar em memória e enviar para download
            virtual_document = BytesIO()
            document.save(virtual_document)
            virtual_document.seek(0)
            
            nome_ficheiro = f"relatorio_{ano}_{mes if tipo_periodo == 'mensal' else 'anual'}.docx"

            return send_file(
                virtual_document,
                as_attachment=True,
                download_name=nome_ficheiro,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    # Lógica GET para exibir a página
    conn, db = get_db()
    db.execute('SELECT id, nome FROM professores ORDER BY nome')
    professores = db.fetchall()
    db.close()
    conn.close()
    ano_atual = datetime.datetime.now().year
    anos = [ano_atual + i for i in range(3)] 
    meses = [(m, datetime.date(2000, m, 1).strftime('%B').capitalize()) for m in range(1, 13)]
    
    return render_template('relatorio.html', anos=anos, meses=meses, professores=professores)

    # --- ROTAS DO CALENDÁRIO ---

@app.route('/calendario/')
@app.route('/calendario/<int:ano>/<int:mes>')
def calendario(ano=None, mes=None):
    agora = datetime.datetime.now()
    if ano is None:
        ano = agora.year
    if mes is None:
        mes = agora.month

    mes_anterior_dt = (datetime.date(ano, mes, 1) - datetime.timedelta(days=1))
    mes_seguinte_dt = (datetime.date(ano, mes, 28) + datetime.timedelta(days=4))

    cal = calendar.Calendar()
    semanas = cal.monthdatescalendar(ano, mes)

    conn, db = get_db()
    db.execute(
    "SELECT id, data, horario, descricao FROM eventos WHERE to_char(data, 'YYYY-MM') = %s ORDER BY horario",
    (f'{ano:04d}-{mes:02d}',)
    )
    eventos_mes = db.fetchall()
    db.close()
    conn.close()

    # --- LÓGICA ATUALIZADA PARA AGRUPAR MÚLTIPLOS EVENTOS POR DIA ---
    eventos_mapa = {}
    for evento in eventos_mes:
        data_evento = evento['data']
        if data_evento not in eventos_mapa:
            eventos_mapa[data_evento] = []
        eventos_mapa[data_evento].append(dict(evento))
    
    meses_nomes = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    nome_mes_atual = meses_nomes[mes - 1]

    return render_template('calendario.html', 
                           semanas=semanas,
                           ano=ano,
                           mes=mes,
                           nome_mes=nome_mes_atual,
                           mes_anterior=mes_anterior_dt,
                           mes_seguinte=mes_seguinte_dt,
                           eventos_mapa=eventos_mapa,
                           hoje=agora.date())
@app.route('/api/eventos/<data>')
def api_get_eventos(data):
    """Retorna todos os eventos de um dia específico em formato JSON."""
    conn, db = get_db()
    db.execute(
    'SELECT id, horario, descricao FROM eventos WHERE data = %s ORDER BY horario', (data,)
    )
    eventos = db.fetchall()
    db.close()
    conn.close()
    # Converte a lista de resultados para JSON
    return jsonify([dict(ix) for ix in eventos])

@app.route('/api/eventos/adicionar', methods=['POST'])
def api_adicionar_evento():
    """Adiciona um novo evento à base de dados."""
    dados = request.get_json()
    data = dados.get('data')
    horario = dados.get('horario')
    descricao = dados.get('descricao')

    if not data or not descricao:
        return jsonify({'status': 'erro', 'mensagem': 'Data e descrição são obrigatórias.'}), 400

    db = get_db()
    db.execute(
        'INSERT INTO eventos (data, horario, descricao) VALUES (%s, %s, %s)',
        (data, horario, descricao)
    )
    db.commit()
    db.close()
    return jsonify({'status': 'sucesso', 'mensagem': 'Evento adicionado!'})

@app.route('/api/eventos/deletar/<int:evento_id>', methods=['POST'])
def api_deletar_evento(evento_id):
    """Apaga um evento específico pelo seu ID."""
    db = get_db()
    db.execute('DELETE FROM eventos WHERE id = %s', (evento_id,))
    db.commit()
    db.close()
    return jsonify({'status': 'sucesso', 'mensagem': 'Evento apagado.'})

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)